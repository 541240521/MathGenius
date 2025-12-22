from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math

class WordExporter:
    def export(self, filepath, questions, with_answers=False, title="数学口算能力测试卷", cols=4, grade_tag="L1 STANDARD", questions_per_page=60):
        doc = Document()
        
        # Set A4 margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        
        # Calculate total pages for questions
        total_q_pages = max(1, (len(questions) + questions_per_page - 1) // questions_per_page)
        
        # Add Questions Pages
        for p in range(total_q_pages):
            if p > 0:
                doc.add_page_break()
            start_idx = p * questions_per_page
            end_idx = min(start_idx + questions_per_page, len(questions))
            page_questions = questions[start_idx:end_idx]
            
            self._add_section(doc, page_questions, title, is_answer=False, cols=cols, grade_tag=grade_tag, 
                              page_num=p+1, total_pages=total_q_pages, start_idx=start_idx, questions_per_page=questions_per_page)
        
        if with_answers:
            total_ans_pages = total_q_pages
            for p in range(total_ans_pages):
                doc.add_page_break()
                start_idx = p * questions_per_page
                end_idx = min(start_idx + questions_per_page, len(questions))
                page_questions = questions[start_idx:end_idx]
                
                self._add_section(doc, page_questions, title + " (参考答案)", is_answer=True, cols=cols, grade_tag=grade_tag, 
                                  page_num=p+1, total_pages=total_ans_pages, start_idx=start_idx, questions_per_page=questions_per_page)
            
        doc.save(filepath)

    def _add_section(self, doc, questions, title, is_answer, cols, grade_tag, page_num, total_pages, start_idx, questions_per_page):
        # Determine dynamic font size based on page capacity
        layout_map = {
            20: {"font": 16, "spacing": 2.0},
            40: {"font": 14, "spacing": 1.5},
            60: {"font": 12, "spacing": 1.15},
            80: {"font": 10, "spacing": 1.0},
            100: {"font": 9, "spacing": 1.0}
        }
        
        conf = layout_map.get(questions_per_page, layout_map[60])
        main_font_size = conf["font"]
        line_spacing = conf["spacing"]
            
        # 1. Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        
        # 2. Info Line
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_text = "姓名：_________  班级：_________  学号：_________  日期：_________  得分：_________"
        run = info_para.add_run(info_text)
        run.font.size = Pt(10)
        
        # 3. Table for Questions
        rows = math.ceil(len(questions) / cols)
        table = doc.add_table(rows=rows, cols=cols)
        table.autofit = True
        
        # Vertical layout mapping for table cells
        for i, q in enumerate(questions):
            col = i // rows
            row = i % rows
            
            if col >= cols: continue # Safety
            
            cell = table.cell(row, col)
            text = q["question"]
            if is_answer:
                display_text = f"{i+1}. {text} {q['answer']}"
            else:
                display_text = f"{i+1}. {text} ____"
                
            cell.text = display_text
            
            # Style cell text
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = line_spacing
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(main_font_size)
                    run.font.bold = True

        # 4. Footer (Simplified for Word)
        # Note: Word footers are usually section-wide, but for simplicity in a practice sheet
        # we can just add a footer paragraph if it's not a complex document.
        # However, to be more robust, we'll set the footer text.
        section = doc.sections[-1]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = f"MATHGENIUS SYSTEM PRO  |  PAGE {page_num} / {total_pages}  |  {grade_tag}"
